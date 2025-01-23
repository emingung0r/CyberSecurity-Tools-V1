import datetime
import os
import importlib
import logging
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from PyPDF2 import PdfReader, PdfWriter

# Log yapılandırması
logging.basicConfig(filename='reports/app.log', level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def check_library(library_name, import_name):
    """Kütüphanenin kurulu olup olmadığını kontrol eder."""
    try:
        importlib.import_module(import_name)
        return True
    except ModuleNotFoundError:
        return False

def generate_report_content(data):
    """Rapor içeriğini biçimlendirir."""
    now = datetime.datetime.now()
    tarih_saat = now.strftime("%Y-%m-%d %H:%M:%S")
    report = f"==================================================\n"
    report += f"Tarih/Saat: {tarih_saat}\n\n"
    report += f">>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>\n"
    report += f"**{data['stage']}: {data.get('stage_name', '')}**\n" #stage_name eklendi
    report += f"**Araç: {data['tool']}**\n"
    report += f"Hedef: {data['target']}**\n"
    report += f"Kullanılan Parametreler: {data['arguments']}\n"
    report += f"<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<\n\n"
    report += f"Sonuç:\n"
    report += f"<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<\n"
    report += str(data['result']) + "\n" # str dönüşümü eklendi
    report += f"==================================================\n"
    return report

def save_report_to_file(report_content, format="txt", filename=None):
    """Raporu dosyaya kaydeder (ekleme modu ile)."""
    if format == "pdf" and not check_library("reportlab", "reportlab.pdfgen"):
        error_message = "1001: PDF raporları oluşturmak için 'reportlab' kütüphanesi kurulu değil. pip install reportlab komutunu kullanın."
        print(error_message)
        logging.error(error_message)
        return error_message
    elif format == "docx" and not check_library("python-docx", "docx"):
        error_message = "1002: Word raporları oluşturmak için 'python-docx' kütüphanesi kurulu değil. pip install python-docx komutunu kullanın."
        print(error_message)
        logging.error(error_message)
        return error_message

    if filename is None:
        filename = f"cybersecurity_report.{format}"
    reports_dir = "reports"
    filepath = os.path.join(reports_dir, filename)
    os.makedirs(reports_dir, exist_ok=True)

    try:
        if format == "txt":
            with open(filepath, "a", encoding="utf-8") as file:
                file.write(report_content)
        elif format == "pdf":
            if os.path.exists(filepath):
                try:
                    reader = PdfReader(filepath)
                    writer = PdfWriter()
                    for page in reader.pages:
                        writer.add_page(page)
                    new_canvas = canvas.Canvas(filepath)
                    textobject = new_canvas.beginText(50, 750)
                    textobject.textLines(report_content)
                    new_canvas.drawText(textobject)
                    writer.add_page(new_canvas)
                    writer.write(filepath)
                except Exception as e:
                    error_message = f"PDF raporuna ekleme sırasında bir hata oluştu: {e}"
                    print(error_message)
                    logging.exception(error_message)
                    return error_message
            else:
                try:
                    c = canvas.Canvas(filepath, pagesize=letter)
                    textobject = c.beginText(50, 750)
                    textobject.textLines(report_content)
                    c.drawText(textobject)
                    c.save()
                except Exception as e:
                    error_message = f"PDF raporu oluşturulurken bir hata oluştu: {e}"
                    print(error_message)
                    logging.exception(error_message)
                    return error_message
        elif format == "docx":
            document = Document()
            if os.path.exists(filepath):
                try:
                    document = Document(filepath)
                    paragraphs = report_content.split("\n\n")
                    for paragraph in paragraphs:
                        document.add_paragraph(paragraph)
                    document.save(filepath)
                except Exception as e:
                    error_message = f"Word raporuna ekleme sırasında bir hata oluştu: {e}"
                    print(error_message)
                    logging.exception(error_message)
                    return error_message

            else:
                try:
                    document = Document()
                    paragraphs = report_content.split("\n\n")
                    for paragraph in paragraphs:
                        document.add_paragraph(paragraph)
                    document.save(filepath)
                except Exception as e:
                    error_message = f"Word raporu oluşturulurken bir hata oluştu: {e}"
                    print(error_message)
                    logging.exception(error_message)
                    return error_message
        print(f"Rapor {filepath} dosyasına kaydedildi.")
        return None
    except PermissionError as e:
        error_message = f"Rapor kaydedilirken izin hatası oluştu: {e}. Rapor klasörüne yazma izniniz olduğundan emin olun."
        print(error_message)
        logging.error(error_message)
        return error_message
    except OSError as e:
        error_message = f"Rapor kaydedilirken disk hatası oluştu: {e}. Diskte yeterli alan olduğundan emin olun."
        print(error_message)
        logging.error(error_message)
        return error_message
    except Exception as e:
        error_message = f"Rapor kaydedilirken beklenmeyen bir hata oluştu: {e}"
        print(error_message)
        logging.exception(error_message)
        return error_message